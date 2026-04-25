"""
DocuNeat - Document Formatter Core
Merapihkan file Word berantakan menjadi laporan profesional
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


MARGIN_PRESETS = {
    'narrow': {'top': 1.27, 'bottom': 1.27, 'left': 1.27,  'right': 1.27},
    'normal': {'top': 2.54, 'bottom': 2.54, 'left': 3.17,  'right': 3.17},
    'wide':   {'top': 2.54, 'bottom': 2.54, 'left': 5.08,  'right': 5.08},
    'mirror': {'top': 2.54, 'bottom': 2.54, 'left': 3.81,  'right': 2.54},
}

LINE_SPACING_MAP = {
    '1.0':  1.0,
    '1.15': 1.15,
    '1.5':  1.5,
    '2.0':  2.0,
}


def detect_heading_level(text):
    t = text.strip()
    if not t or len(t) > 120:
        return 0

    # BAB I, BAB II, BAB 1
    if re.match(r'^bab\s+([ivxlcdm]+|\d+)\b', t, re.IGNORECASE):
        return 1

    # Nama seksi umum laporan akademik
    SECTIONS_H1 = {
        'pendahuluan', 'latar belakang', 'rumusan masalah',
        'tujuan penelitian', 'tujuan', 'manfaat penelitian', 'manfaat',
        'kajian pustaka', 'tinjauan pustaka', 'landasan teori',
        'metodologi penelitian', 'metodologi', 'metode penelitian',
        'hasil dan pembahasan', 'hasil penelitian', 'pembahasan',
        'kesimpulan', 'kesimpulan dan saran', 'saran',
        'daftar pustaka', 'daftar referensi', 'referensi',
        'lampiran', 'abstrak', 'abstract',
        'kata pengantar', 'daftar isi', 'daftar gambar', 'daftar tabel',
        'penutup', 'analisis', 'tinjauan',
    }
    if t.lower().rstrip('.,:') in SECTIONS_H1:
        return 1

    # ALL CAPS tanpa angka, panjang 4-80 char
    if t.isupper() and 4 <= len(t) <= 80 and not re.search(r'\d', t) and not t.endswith(':'):
        return 1

    # 1.1.1 Judul
    if re.match(r'^\d+\.\d+\.\d+\s+\S', t):
        return 3

    # 1.1 Judul
    if re.match(r'^\d+\.\d+\s+\S', t):
        return 2

    # I. JUDUL
    if re.match(r'^[IVXLCDM]+\.\s+[A-Z]', t):
        return 1

    return 0


def should_keep_left(para, text):
    """
    True = jangan di-justify, biarkan rata kiri.
    Berlaku untuk: list, keterangan gambar/tabel, label formulir,
    teks pendek, tanda tangan, font monospace.
    """
    # Ada numbering XML (bullet/numbered list)
    if para._p.find(qn('w:numPr')) is not None:
        return True

    # Bullet character manual
    if re.match(r'^[\u2022\u2023\u25e6\u2043\u2219\u25aa\u25ab\u25cf\u25cb\u25a0\u25a1•·▪▸►-]\s', text):
        return True

    # Keterangan gambar, tabel, grafik, dst
    if re.match(r'^(gambar|tabel|grafik|bagan|diagram|foto|ilustrasi|lampiran|figure|table)\s*[\d\.]', text, re.IGNORECASE):
        return True

    # Label formulir: "Nama :", "NIM    :"
    if re.search(r'\s*:\s*$', text):
        return True
    if re.match(r'^[A-Za-z /,.\(\)]{2,40}\s{0,6}:', text):
        return True

    # Teks hanya 1-3 kata
    if len(text.split()) <= 3:
        return True

    # Kota + tanggal: "Jakarta, 1 Januari 2024"
    if re.match(r'^[A-Za-z ]+,\s+\d{1,2}\s+[A-Za-z]+\s+\d{4}', text):
        return True

    # Font monospace pada run pertama
    if para.runs:
        fn = (para.runs[0].font.name or '').lower()
        if any(m in fn for m in ('courier', 'consolas', 'mono', 'lucida console', 'fixedsys')):
            return True

    return False


def set_run_font(run, font_name, font_size_pt):
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'),    font_name)
    rFonts.set(qn('w:hAnsi'),   font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'),      font_name)


def set_paragraph_spacing(para, line_spacing_val, space_before=0, space_after=6):
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    twips = {1.0: '240', 1.15: '276', 1.5: '360', 2.0: '480'}
    spacing.set(qn('w:line'),     twips.get(line_spacing_val, '360'))
    spacing.set(qn('w:lineRule'), 'auto')
    spacing.set(qn('w:before'),   str(int(space_before * 20)))
    spacing.set(qn('w:after'),    str(int(space_after  * 20)))


def _make_page_field_run(font_name):
    """Buat w:r berisi field PAGE."""
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '20')
    rPr.append(sz)
    run.append(rPr)

    fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin')
    run.append(fc1)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve'); instr.text = ' PAGE '
    run.append(instr)
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'end')
    run.append(fc2)
    return run


def add_page_numbers(doc, position='bottom_center', font_name='Times New Roman'):
    """
    position: 'bottom_center' | 'bottom_right' | 'top_right'
    """
    align_map = {
        'bottom_center': WD_ALIGN_PARAGRAPH.CENTER,
        'bottom_right':  WD_ALIGN_PARAGRAPH.RIGHT,
        'top_right':     WD_ALIGN_PARAGRAPH.RIGHT,
    }
    align = align_map.get(position, WD_ALIGN_PARAGRAPH.CENTER)
    use_header = (position == 'top_right')

    for section in doc.sections:
        section.different_first_page_header_footer = False
        container = section.header if use_header else section.footer
        container.is_linked_to_previous = False

        for p in container.paragraphs:
            p.clear()

        para = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
        para.alignment = align

        # Hapus run lama
        for r in para._p.findall(qn('w:r')):
            para._p.remove(r)
        para._p.append(_make_page_field_run(font_name))


def format_document(input_path, output_path, options):
    doc = Document(input_path)

    font_name     = options.get('font',            'Times New Roman')
    font_size     = int(options.get('font_size',   12))
    line_spacing  = LINE_SPACING_MAP.get(options.get('line_spacing', '1.5'), 1.5)
    margin_name   = options.get('margin',          'normal')
    fix_headings  = options.get('fix_headings',    True)
    fix_spacing   = options.get('fix_spacing',     True)
    fix_alignment = options.get('fix_alignment',   True)
    add_pn        = options.get('add_page_numbers', False)
    pn_position   = options.get('page_number_position', 'bottom_center')
    remove_blank  = options.get('remove_extra_blank', True)

    margins = MARGIN_PRESETS.get(margin_name, MARGIN_PRESETS['normal'])

    stats = {
        'paragraphs_formatted':     0,
        'headings_detected':        0,
        'blank_paragraphs_removed': 0,
        'kept_left_align':          0,
        'font_changed':             font_name,
        'font_size':                font_size,
    }

    # Margins
    for section in doc.sections:
        section.top_margin    = Cm(margins['top'])
        section.bottom_margin = Cm(margins['bottom'])
        section.left_margin   = Cm(margins['left'])
        section.right_margin  = Cm(margins['right'])

    # Hapus baris kosong berlebih
    if remove_blank:
        consecutive = 0
        to_remove = []
        for para in doc.paragraphs:
            if not para.text.strip():
                consecutive += 1
                if consecutive > 1:
                    to_remove.append(para._p)
                    stats['blank_paragraphs_removed'] += 1
            else:
                consecutive = 0
        for p in to_remove:
            parent = p.getparent()
            if parent is not None:
                parent.remove(p)

    # Format setiap paragraf
    for para in doc.paragraphs:
        text = para.text.strip()

        heading_level = 0
        if fix_headings and text:
            heading_level = detect_heading_level(text)

        if heading_level > 0:
            # Heading
            try:
                para.style = doc.styles[f'Heading {heading_level}']
            except Exception:
                pass

            h_size = font_size + (2 if heading_level == 1 else 1 if heading_level == 2 else 0)
            for run in para.runs:
                set_run_font(run, font_name, h_size)
                run.font.bold = True

            if fix_alignment:
                para.alignment = (WD_ALIGN_PARAGRAPH.CENTER if heading_level == 1
                                  else WD_ALIGN_PARAGRAPH.LEFT)

            set_paragraph_spacing(para, line_spacing, space_before=12, space_after=6)
            stats['headings_detected'] += 1

        else:
            # Paragraf biasa
            for run in para.runs:
                set_run_font(run, font_name, font_size)

            if fix_alignment and text:
                if should_keep_left(para, text):
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    stats['kept_left_align'] += 1
                elif para.alignment not in (WD_ALIGN_PARAGRAPH.CENTER,
                                             WD_ALIGN_PARAGRAPH.RIGHT):
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if fix_spacing:
                set_paragraph_spacing(para, line_spacing, space_before=0, space_after=6)

            if text:
                stats['paragraphs_formatted'] += 1

    # Format tabel — sel rata kiri, bukan justify
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_run_font(run, font_name, font_size - 1)
                    if fix_alignment and para.text.strip():
                        if para.alignment not in (WD_ALIGN_PARAGRAPH.CENTER,
                                                   WD_ALIGN_PARAGRAPH.RIGHT):
                            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Nomor halaman
    if add_pn:
        add_page_numbers(doc, position=pn_position, font_name=font_name)

    doc.save(output_path)
    return stats
